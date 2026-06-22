"""Markdown document artifact helpers."""

from __future__ import annotations

import os
import re
import uuid
from pathlib import Path

DEFAULT_ARTIFACT_ROOT = Path(__file__).resolve().parents[2] / "artifacts"
DEFAULT_REVIEW_DOC_ARTIFACT_DIR = DEFAULT_ARTIFACT_ROOT / "review_docs"
DEFAULT_EXERCISE_ARTIFACT_DIR = DEFAULT_ARTIFACT_ROOT / "exercises"
DEFAULT_CODE_PRACTICE_ARTIFACT_DIR = DEFAULT_ARTIFACT_ROOT / "code-practice"

_ARTIFACT_KIND_CONFIG = {
    "review_docs": {
        "env_var": "REVIEW_DOC_ARTIFACT_DIR",
        "default_dir": DEFAULT_REVIEW_DOC_ARTIFACT_DIR,
        "url_prefix": "/artifacts/review-docs",
        "filename_default": "review-doc",
    },
    "exercises": {
        "env_var": "EXERCISE_ARTIFACT_DIR",
        "default_dir": DEFAULT_EXERCISE_ARTIFACT_DIR,
        "url_prefix": "/artifacts/exercises",
        "filename_default": "exercises",
    },
    "code_practice": {
        "env_var": "CODE_PRACTICE_ARTIFACT_DIR",
        "default_dir": DEFAULT_CODE_PRACTICE_ARTIFACT_DIR,
        "url_prefix": "/artifacts/code-practice",
        "filename_default": "code-practice",
    },
}


def get_review_doc_artifact_dir() -> Path:
    """Return the directory used for generated Markdown review documents."""
    root = Path(os.getenv("REVIEW_DOC_ARTIFACT_DIR", str(DEFAULT_REVIEW_DOC_ARTIFACT_DIR)))
    root.mkdir(parents=True, exist_ok=True)
    return root.resolve()


def get_exercise_artifact_dir() -> Path:
    """Return the directory used for generated exercise documents."""
    root = Path(os.getenv("EXERCISE_ARTIFACT_DIR", str(DEFAULT_EXERCISE_ARTIFACT_DIR)))
    root.mkdir(parents=True, exist_ok=True)
    return root.resolve()


def get_code_practice_artifact_dir() -> Path:
    """Return the directory used for generated code-practice documents."""
    root = Path(os.getenv("CODE_PRACTICE_ARTIFACT_DIR", str(DEFAULT_CODE_PRACTICE_ARTIFACT_DIR)))
    root.mkdir(parents=True, exist_ok=True)
    return root.resolve()


def _get_document_artifact_dir(artifact_kind: str) -> Path:
    config = _ARTIFACT_KIND_CONFIG.get(artifact_kind)
    if not config:
        raise ValueError(f"Unsupported document artifact kind: {artifact_kind}")
    root = Path(os.getenv(str(config["env_var"]), str(config["default_dir"])))
    root.mkdir(parents=True, exist_ok=True)
    return root.resolve()


def _safe_filename_stem(value: str, default: str = "review-doc") -> str:
    cleaned = re.sub(r"[^\w\u4e00-\u9fff.-]+", "-", value.strip(), flags=re.UNICODE)
    cleaned = cleaned.strip(".-")[:80]
    return cleaned or default


def _extract_first_python_code_block(markdown_text: str) -> str:
    match = re.search(
        r"(?s)```(?:python|py)\s*\n(?P<code>.+?)```",
        markdown_text or "",
    )
    return match.group("code").strip() + "\n" if match else ""


def _is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _is_table_separator(line: str) -> bool:
    cells = _parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _clean_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def _add_markdown_table(document, rows: list[list[str]]) -> None:
    rows = [row for row in rows if row]
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(cols):
            value = row[col_idx] if col_idx < len(row) else ""
            table.cell(row_idx, col_idx).text = _clean_inline_markdown(value)


def _write_docx_artifact(markdown_text: str, title: str, file_path: Path) -> None:
    from docx import Document

    document = Document()
    lines = markdown_text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if _is_table_row(stripped):
            table_rows: list[list[str]] = []
            while index < len(lines) and _is_table_row(lines[index]):
                if not _is_table_separator(lines[index]):
                    table_rows.append(_parse_table_row(lines[index]))
                index += 1
            _add_markdown_table(document, table_rows)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            document.add_heading(_clean_inline_markdown(heading.group(2)), level=level)
            index += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet:
            document.add_paragraph(_clean_inline_markdown(bullet.group(1)), style="List Bullet")
            index += 1
            continue

        document.add_paragraph(_clean_inline_markdown(stripped))
        index += 1

    if not lines:
        document.add_heading(title, level=1)
    document.save(file_path)


def create_document_artifact(
    markdown_text: str,
    title: str,
    artifact_kind: str = "review_docs",
) -> dict:
    """Save Markdown text as .md and .docx artifacts and return public metadata."""
    config = _ARTIFACT_KIND_CONFIG.get(artifact_kind)
    if not config:
        raise ValueError(f"Unsupported document artifact kind: {artifact_kind}")

    artifact_id = uuid.uuid4().hex
    filename_stem = _safe_filename_stem(title, default=str(config["filename_default"]))
    filename = f"{filename_stem}.md"
    docx_filename = f"{filename_stem}.docx"
    artifact_dir = _get_document_artifact_dir(artifact_kind) / artifact_id
    artifact_dir.mkdir(parents=True, exist_ok=True)
    file_path = artifact_dir / filename
    docx_path = artifact_dir / docx_filename
    file_path.write_text(markdown_text.rstrip() + "\n", encoding="utf-8")
    _write_docx_artifact(markdown_text, title, docx_path)
    url_prefix = str(config["url_prefix"])

    return {
        "artifact_id": artifact_id,
        "filename": filename,
        "docx_filename": docx_filename,
        "markdown_url": f"{url_prefix}/{artifact_id}/{filename}",
        "docx_url": f"{url_prefix}/{artifact_id}/{docx_filename}",
        "title": title,
    }


def create_markdown_artifact(markdown_text: str, title: str) -> dict:
    """Save review-doc Markdown text as .md and .docx artifacts."""
    return create_document_artifact(markdown_text, title, artifact_kind="review_docs")


def create_code_practice_artifact(
    markdown_text: str,
    title: str,
    python_code: str | None = None,
) -> dict:
    """Save code-practice Markdown as .md/.docx plus an extracted .py file."""
    artifact = create_document_artifact(
        markdown_text=markdown_text,
        title=title,
        artifact_kind="code_practice",
    )
    code = str(python_code or "").strip()
    if not code:
        code = _extract_first_python_code_block(markdown_text).strip()
    if not code:
        code = 'print("请在 Markdown 文档中查看代码实操内容")'

    artifact_id = str(artifact["artifact_id"])
    python_filename = Path(str(artifact["filename"])).with_suffix(".py").name
    artifact_dir = get_code_practice_artifact_dir() / artifact_id
    artifact_dir.mkdir(parents=True, exist_ok=True)
    (artifact_dir / python_filename).write_text(code.rstrip() + "\n", encoding="utf-8")

    return {
        **artifact,
        "python_filename": python_filename,
        "python_url": f"/artifacts/code-practice/{artifact_id}/{python_filename}",
        "markdown": markdown_text,
    }
